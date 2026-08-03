#!/usr/bin/env python3
"""Transcript + Flow prompts + five reference images per shot, as a Google Doc.

Three things per segment of the locked audio:
  1. the exact words spoken
  2. a prompt you can paste straight into Flow, with nothing to edit out
  3. five reference images, so the look is chosen rather than accepted

The one structural rule the document exists to fix: WHEN JIMMY TALKS, YOU SEE
JIMMY. A bite is his own recorded voice at a known timecode; the picture over
it has to be the real man saying the real words. Generating a likeness of a
real person delivering a real quote would be a fabrication, and it is also
what made the last cut feel out of sync with the audio. Those segments get a
source and an in-point instead of a prompt.

WHY THIS BUILDS A .DOCX AND NOT HTML
------------------------------------
The previous version wrote HTML with <img src="https://..."> and let Drive
fetch each image at conversion time. That works only for hosts that permit
hotlinking, which in practice meant the references had to come from Pexels and
Wikimedia - and being stuck with those two libraries is a large part of why
the references were poor. A .docx carries the image bytes inside the file, so
the pictures can come from anywhere, nothing is fetched at conversion time,
and no image can silently fail to appear.

The document is still updated in place by title, so the URL never changes.

Run:  py -3.12 pipeline/shotlist_doc.py            (uses cached references)
      py -3.12 pipeline/shotlist_doc.py --refresh  (re-runs every search)
      py -3.12 pipeline/shotlist_doc.py --refresh 15 16   (just those segments)
      py -3.12 pipeline/shotlist_doc.py --local    (build the .docx, no upload)
"""

from __future__ import annotations

import json
import pickle
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "pipeline"))

import imgref                                                    # noqa: E402
from shot_prompts import (MED_LOCAL, PROMPTS, REAL,               # noqa: E402
                          REF_QUERY)

EDL = ROOT / "manifest/edl_full.json"
REFDIR = ROOT / "work/refs"
OUT_DOCX = ROOT / "final_video/MRBEAST_SHOTLIST.docx"
TOKEN = Path(r"C:\Users\avina\OneDrive\Desktop\Claude Projects\XL Eagle"
             r"\Tools and Dashboards\Pickle creation\tokens"
             r"\token_slides_info_xleagle_com.pickle")
TITLE = "MrBeast — transcript, shot prompts + references"
DOC_MIME = "application/vnd.google-apps.document"
DOCX_MIME = ("application/vnd.openxmlformats-officedocument"
             ".wordprocessingml.document")
N_REFS = 5

SOURCE_NAME = {
    "cLRLEnPaJLM": "Joe Rogan Experience #1788",
    "FjrJ2DJN_pA": "The Diary Of A CEO",
    "9IQ_ldV9z_A": "Colin and Samir (June 2023)",
    "c8VcUnz3nVc": "Colin and Samir — The Full Story of MrBeast",
    "7r3ORKgNUjw": "Airrack — My 600 Day Transformation",
}
CHAPTER_TITLE = {
    "open": "COLD OPEN", "origin": "1 — THE ATHLETE, AND WHAT TOOK IT",
    "illness": "2 — WHAT THE DISEASE ACTUALLY DOES",
    "machine": "3 — THE MACHINE HE BUILT INSTEAD", "contract": "4 — THE PACT",
    "protocol": "5 — WHAT HE ACTUALLY DOES",
    "limit": "6 — WHERE IT STOPS WORKING",
    "fall": "7 — THE MACHINE TAKES IT BACK",
    "resolution": "8 — WHAT HE GOT TO KEEP",
}


# ---------------------------------------------------------------- references

def resolve_refs(only: set | None = None, refresh: bool = False) -> dict:
    """{segment key: [image path, ...]} - cached on disk, five per segment.

    A medical segment gets its already-licensed local stills first, then the
    search fills the row up to five. That guarantees every medical block has
    at least one image that is verified real Crohn's material.
    """
    REFDIR.mkdir(parents=True, exist_ok=True)
    manifest_path = REFDIR / "manifest.json"
    manifest = {}
    if manifest_path.exists():
        try:
            manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            manifest = {}

    todo = {}
    for seg, (query, kind) in REF_QUERY.items():
        key = str(seg)
        if only is not None and seg not in only:
            continue
        have = manifest.get(key, {})
        if have.get("query") == query and len(have.get("files", [])) >= N_REFS \
                and not refresh:
            continue
        todo[key] = (query, kind)

    if todo:
        print(f"[search] {len(todo)} segment(s)", flush=True)
        results = imgref.search_many(todo)
        for key, recs in results.items():
            seg = float(key) if "." in key else int(key)
            outdir = REFDIR / key.replace("-", "m").replace(".", "_")
            outdir.mkdir(parents=True, exist_ok=True)
            for old in outdir.glob("*.jpg"):
                old.unlink()

            picked, provenance = [], []
            for name in MED_LOCAL.get(seg, []):                  # pinned first
                im = imgref.local(name)
                if im is not None:
                    p = outdir / f"{len(picked)}.jpg"
                    im.save(p, quality=88)
                    picked.append(str(p))
                    provenance.append({"source": "dossier/mrbeast/medical",
                                       "file": name})
            for rec in recs:
                if len(picked) >= N_REFS:
                    break
                im = imgref.fetch(rec)
                if im is None:
                    continue
                p = outdir / f"{len(picked)}.jpg"
                im.save(p, quality=88)
                picked.append(str(p))
                provenance.append({"source": rec.get("page"),
                                   "title": rec.get("title")})

            manifest[key] = {"query": REF_QUERY[seg][0], "files": picked,
                             "provenance": provenance}
            print(f"  [{key:>5}] kept {len(picked)}/{N_REFS}", flush=True)
        manifest_path.write_text(json.dumps(manifest, indent=1),
                                 encoding="utf-8")

    return {k: v.get("files", []) for k, v in manifest.items()}


def sheets(refs: dict) -> None:
    """Contact sheets of everything chosen, for an eyes-on pass."""
    from PIL import Image
    out = REFDIR / "_sheets"
    out.mkdir(parents=True, exist_ok=True)
    keys = sorted(refs, key=lambda k: float(k))
    for start in range(0, len(keys), 8):
        cells = []
        for k in keys[start:start + 8]:
            for i in range(N_REFS):
                files = refs[k]
                im = None
                if i < len(files) and Path(files[i]).exists():
                    im = Image.open(files[i]).convert("RGB")
                cells.append((im, f"seg {k} · {i + 1}"))
        p = out / f"sheet_{start // 8:02d}.png"
        imgref.contact_sheet(cells, p)
        print(f"  [sheet] {p}")


# ---------------------------------------------------------------------- docx

def _grey(run, size=8):
    from docx.shared import Pt, RGBColor
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


def block(doc, label, stamp, spoken, prompt, images, real=None):
    from docx.shared import Inches, Pt

    p = doc.add_paragraph()
    p.add_run(label).bold = True
    _grey(p.add_run("   " + stamp))

    if spoken:
        doc.add_paragraph(spoken)

    if real:
        src, inpoint, note = real
        p = doc.add_paragraph()
        p.add_run("USE THE REAL FOOTAGE. ").bold = True
        p.add_run(f"{src}, around {inpoint} — {note}")

    if prompt:
        p = doc.add_paragraph()
        r = p.add_run("FLOW PROMPT: " + prompt)
        r.italic = True

    usable = [f for f in (images or [])[:N_REFS] if Path(f).exists()]
    if usable:
        t = doc.add_table(rows=2, cols=N_REFS)
        for i in range(N_REFS):
            cell = t.cell(0, i)
            cell.paragraphs[0].alignment = 1
            if i < len(usable):
                cell.paragraphs[0].add_run().add_picture(
                    usable[i], width=Inches(1.15))
            n = t.cell(1, i).paragraphs[0]
            n.alignment = 1
            if i < len(usable):
                r = n.add_run(str(i + 1))
                r.font.size = Pt(8)
        p = doc.add_paragraph()
        _grey(p.add_run("references only — delete the four you do not want"))

    doc.add_paragraph()


def tc(t: float) -> str:
    return f"{int(t // 60)}:{t % 60:05.2f}"


def build_docx(refs: dict) -> Path:
    import docx

    edl = json.loads(EDL.read_text(encoding="utf-8"))
    doc = docx.Document()

    doc.add_heading("The Disease That Built MrBeast", level=0)
    doc.add_heading("Transcript, shot prompts and references", level=3)
    p = doc.add_paragraph()
    p.add_run(
        "The audio is locked and approved. Each block below is one segment of "
        "it, in order: the exact words, a prompt you can paste straight into "
        "Flow, and five reference images for the look. The references are "
        "search results and licensed medical stills, not frames of the film."
    ).italic = True

    doc.add_heading("The one rule", level=2)
    p = doc.add_paragraph()
    p.add_run("When Jimmy talks, you see Jimmy.").bold = True
    p.add_run(
        " A BITE is his own recorded voice at a known timecode, so the picture "
        "over it must be the real man saying the real words — pulled from the "
        "real interview, never generated and never a cutaway. That is what "
        "made the last cut feel out of sync. NARRATION is the narrator, and "
        "that is where the generated and stock material belongs.")
    doc.add_paragraph(
        "There are eight places where his voice plays without his face, and "
        "every one says why: six are Airrack's video, where no Jimmy-only "
        "frame exists, and two are a Rogan stretch with a web page burned "
        "into the shot behind him. Those eight now carry a prompt, because "
        "something has to be on screen.")

    doc.add_heading("How to read the prompts", level=2)
    doc.add_paragraph(
        "Each prompt is keyed to the sentence above it. Where the line states "
        "something specific — three hundred million subscribers, the "
        "inflammation goes through the whole thickness of the wall, six "
        "hundred days — the shot shows that thing. Where the line is "
        "interpretation rather than fact, the shot is a metaphor, but one "
        "that resolves in about a second. No shot is cast as Jimmy: no "
        "stand-in body, no silhouette training, no legs walking.")

    chapter = None
    for s in edl["segs"]:
        i = s["i"]
        if s["chapter"] != chapter:
            chapter = s["chapter"]
            doc.add_heading(CHAPTER_TITLE.get(chapter, chapter), level=2)
        if i == 0:
            block(doc, "LEAD-IN — music only", "0:00–0:02", "",
                  PROMPTS.get(-1, ""), refs.get("-1", []))

        kind, text = s["kind"], (s.get("text") or "").strip()
        stamp = f"{tc(s['start'])}–{tc(s['end'])}"
        if kind == "beat":
            label, spoken = "BEAT — music only", ""
        elif kind == "bite":
            src = SOURCE_NAME.get(s.get("source", ""), s.get("source", ""))
            label, spoken = f"BITE — HIS OWN VOICE · {src}", f"“{text}”"
        elif kind == "card":
            label, spoken = "CARD", text
        else:
            label, spoken = "NARRATION", text

        block(doc, label, stamp, spoken, PROMPTS.get(i, ""),
              refs.get(str(i), []), REAL.get(i))

        if i == 6:
            block(doc, "TITLE BREAK — 4s, music only", "0:54.10–0:58.10", "",
                  PROMPTS.get(6.5, ""), refs.get("6.5", []))

    block(doc, "TAIL — music only", "12:16–12:19", "",
          "The end card holds on black. Nothing else on screen.", [])

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    return OUT_DOCX


# -------------------------------------------------------------------- upload

def creds():
    c = pickle.load(open(TOKEN, "rb"))
    if not c.valid and c.expired and c.refresh_token:
        from google.auth.transport.requests import Request
        c.refresh(Request())
        pickle.dump(c, open(TOKEN, "wb"))
    return c


def upload(path: Path) -> dict:
    from googleapiclient.discovery import build
    from googleapiclient.http import MediaFileUpload

    c = creds()
    drive = build("drive", "v3", credentials=c, cache_discovery=False)
    q = f"name = '{TITLE}' and mimeType = '{DOC_MIME}' and trashed = false"
    found = drive.files().list(q=q, fields="files(id)").execute().get("files", [])
    media = MediaFileUpload(str(path), mimetype=DOCX_MIME, resumable=True)
    if found:
        f = drive.files().update(fileId=found[0]["id"], media_body=media,
                                 fields="id,name,webViewLink").execute()
        print("[drive] updated in place — same URL")
    else:
        f = drive.files().create(
            body={"name": TITLE, "mimeType": DOC_MIME}, media_body=media,
            fields="id,name,webViewLink").execute()
        print("[drive] created")

    docs = build("docs", "v1", credentials=c, cache_discovery=False)
    d = docs.documents().get(documentId=f["id"]).execute()
    f["n_images"] = len(d.get("inlineObjects") or {})
    return f


def main() -> int:
    args = sys.argv[1:]
    local_only = "--local" in args
    refresh = "--refresh" in args
    only = None
    if refresh:
        nums = [a for a in args if not a.startswith("--")]
        if nums:
            only = {float(n) if "." in n else int(n) for n in nums}

    refs = resolve_refs(only=only, refresh=refresh)
    got = sum(1 for v in refs.values() if v)
    print(f"[refs] {got}/{len(REF_QUERY)} segments have references")

    if "--sheets" in args:
        sheets(refs)

    path = build_docx(refs)
    print(f"[docx] {path}  {path.stat().st_size / 1024:.0f} KB")

    if local_only:
        return 0

    f = upload(path)
    print(f"\n[OK] {f['name']}")
    print(f"     {f['n_images']} reference images embedded")
    print(f"     {f['webViewLink']}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
